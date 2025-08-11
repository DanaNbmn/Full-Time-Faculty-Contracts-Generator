
import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime

st.set_page_config(page_title="ADU Faculty Contract Generator", page_icon="📄", layout="centered")

# ========= 1) CONFIG =========
DEFAULT_TEMPLATE_PATH = "Faculty_Offer_Letter_Template_BenefitsBlock.docx"
DATE_FORMAT = "%d %B %Y"  # e.g., 11 August 2025

# ========= 2) BENEFITS RULES =========
BENEFITS = {
    "_shared": {
        "children_school_allowance": {"AD/Dubai": 60000, "AA": 50000},
        "tuition_waiver": {"employee": 75, "dependent": 50, "immediate_family": 25},
        "annual_ticket": "1+1+2 Economy",
        "repatriation_ticket": "1+1+2 Economy",
        "health_insurance": "1+1+3",
    },
    "Professor": {
        "annual_leave_days": 56, "relocation_allowance": 3000, "repatriation_allowance": 3000,
        "joining_ticket_international": "1+1+2 Economy",
        "housing_allowance_k": {"AD/Dubai": {"Single": 45, "Married": 60}, "AA": {"Single": 35, "Married": 45}},
        "furniture_allowance_k_once": {"AD/Dubai": {"Single": 20, "Married": 30}, "AA": {"Single": 20, "Married": 30}},
    },
    "Associate / Sr. Lecturer": {
        "annual_leave_days": 56, "relocation_allowance": 3000, "repatriation_allowance": 3000,
        "joining_ticket_international": "1+1+2 Economy",
        "housing_allowance_k": {"AD/Dubai": {"Single": 45, "Married": 60}, "AA": {"Single": 35, "Married": 45}},
        "furniture_allowance_k_once": {"AD/Dubai": {"Single": 20, "Married": 30}, "AA": {"Single": 20, "Married": 30}},
    },
    "Assistant / Lecturer": {
        "annual_leave_days": 56, "relocation_allowance": 3000, "repatriation_allowance": 3000,
        "joining_ticket_international": "1+1+2 Economy",
        "housing_allowance_k": {"AD/Dubai": {"Single": 45, "Married": 60}, "AA": {"Single": 35, "Married": 45}},
        "furniture_allowance_k_once": {"AD/Dubai": {"Single": 20, "Married": 30}, "AA": {"Single": 20, "Married": 30}},
    },
    "Senior Instructor": {
        "annual_leave_days": 42, "relocation_allowance": 3000, "repatriation_allowance": 2000,
        "joining_ticket_international": "1+1+2 Economy",
        "housing_allowance_k": {"AD/Dubai": {"Single": 35, "Married": 45}, "AA": {"Single": 30, "Married": 40}},
        "furniture_allowance_k_once": {"AD/Dubai": {"Single": 12, "Married": 15}, "AA": {"Single": 12, "Married": 15}},
    },
    "Instructor": {
        "annual_leave_days": 42, "relocation_allowance": 3000, "repatriation_allowance": 2000,
        "joining_ticket_international": "1+1+2 Economy",
        "housing_allowance_k": {"AD/Dubai": {"Single": 35, "Married": 45}, "AA": {"Single": 30, "Married": 40}},
        "furniture_allowance_k_once": {"AD/Dubai": {"Single": 12, "Married": 15}, "AA": {"Single": 12, "Married": 15}},
    },
}

# ========= 3) HELPERS =========
def campus_key(campus: str) -> str:
    # Dubai follows Abu Dhabi rules
    return "AD/Dubai" if campus in ["Abu Dhabi", "Dubai", "AD/Dubai"] else "AA"

def fmt_amt(n: int) -> str:
    return f"{int(n):,}"

def compute_benefits_mapping(rank: str, marital: str, campus: str, is_international: bool):
    R = BENEFITS[rank]
    S = BENEFITS["_shared"]
    ckey = campus_key(campus)

    housing = R["housing_allowance_k"][ckey][marital] * 1000
    furniture = R["furniture_allowance_k_once"][ckey][marital] * 1000

    # Commencement (joining) ticket only for international hires
    joining_value = R["joining_ticket_international"] if is_international else ""

    edu_per_child = S["children_school_allowance"][ckey]
    edu_total = S["children_school_allowance"][ckey]

    return {
        "JOINING_TICKET": joining_value,
        "HOUSING_ALLOWANCE": fmt_amt(housing),
        "FURNITURE_ALLOWANCE": fmt_amt(furniture),
        "EDUCATION_ALLOWANCE_PER_CHILD": fmt_amt(edu_per_child),
        "EDUCATION_ALLOWANCE_TOTAL": fmt_amt(edu_total),
        "TUITION_EMPLOYEE": S["tuition_waiver"]["employee"],
        "TUITION_DEPENDENT": S["tuition_waiver"]["dependent"],
        "TUITION_IMMEDIATE": S["tuition_waiver"]["immediate_family"],
        "RELOCATION_ALLOWANCE": fmt_amt(R["relocation_allowance"]),
        "REPARIATION_ALLOWANCE": fmt_amt(R["repatriation_allowance"]),
        "REPATRIATION_TICKET": S["repatriation_ticket"],     # always applies
        "HEALTH_INSURANCE": S["health_insurance"],           # always applies
        "ANNUAL_LEAVE_DAYS": R["annual_leave_days"],
    }

def compose_benefits_block(m: dict) -> str:
    join_line = f"- Commencement Air Tickets: {m['JOINING_TICKET']}\n" if m.get("JOINING_TICKET") else ""
    return (
        f"- Accommodation: Unfurnished on-campus accommodation based on availability, or a housing allowance of AED {m['HOUSING_ALLOWANCE']} per year (paid monthly) will be provided based on eligibility.\n"
        f"- Furniture Allowance: AED {m['FURNITURE_ALLOWANCE']} provided at the commencement of employment as a forgivable loan amortized over three (3) years. Should you leave ADU before completing three years of service, the amount will be repayable on a pro-rata basis.\n"
        "- Annual Leave Airfare: Cash in lieu of economy class air tickets for yourself, your spouse, and up to two (2) eligible dependent children under the age of 21 years residing in the UAE, based on ADU’s published schedule of rates including your country of origin. This amount will be paid annually in the month of May, prorated to your joining date.\n"
        f"{join_line}"
        f"- Relocation Allowance: Up to AED {m['RELOCATION_ALLOWANCE']} at the commencement of employment to support the relocation of personal effects to ADU-provided accommodation. Reimbursement will be subject to submission of original receipts.\n"
        f"- Repatriation Air Tickets: {m['REPATRIATION_TICKET']}\n"
        f"- Repatriation Allowance: AED {m['REPARIATION_ALLOWANCE']} upon conclusion of your contract, applicable only upon completion of two (2) years of continuous service with ADU.\n"
        f"- Medical Insurance: {m['HEALTH_INSURANCE']}\n"
        f"- Annual Leave Entitlement: {m['ANNUAL_LEAVE_DAYS']} calendar days of paid annual leave.\n"
        f"- School Fee Subsidy: An annual subsidy of AED {m['EDUCATION_ALLOWANCE_PER_CHILD']} per eligible child under the age of 21 years residing in the UAE under your sponsorship, up to a maximum of AED {m['EDUCATION_ALLOWANCE_TOTAL']} per family. This benefit applies only to married individuals with children under their sponsorship.\n"
        f"- ADU Tuition Waiver: {m['TUITION_EMPLOYEE']}% tuition fee deduction for yourself, {m['TUITION_DEPENDENT']}% for dependents, and {m['TUITION_IMMEDIATE']}% for immediate family members, in accordance with ADU policy. This benefit is applicable upon completion of one (1) year of service."
    )

def replace_placeholders(doc: Document, mapping: dict):
    def replace_in_paragraph(par):
        text = par.text
        for k, v in mapping.items():
            token = f"{{{{{k}}}}}"
            if token in text:
                text = text.replace(token, str(v))
        for _ in range(len(par.runs)):
            par.runs[0].text = ""
            del par.runs[0]
        par.add_run(text)

    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

def generate_docx(template_bytes: bytes | None, mapping: dict) -> bytes:
    doc = Document(BytesIO(template_bytes)) if template_bytes else Document(DEFAULT_TEMPLATE_PATH)
    replace_placeholders(doc, mapping)
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()

# ========= 4) UI =========
st.title("📄 ADU – Faculty Offer Letter Generator")

with st.form("offer_form", clear_on_submit=False):
    st.subheader("Candidate & Position")
    c1, c2 = st.columns(2)
    with c1:
        candidate_id = st.text_input("ID (Ref)", placeholder="TEG/2025/001")
        salutation = st.selectbox("Salutation", ["Dr.", "Mr.", "Ms.", "Prof.", "Eng."], index=0)
        candidate_name = st.text_input("Candidate Name", placeholder="Full Name")
        telephone = st.text_input("Telephone", placeholder="+971-XX-XXX-XXXX")
        personal_email = st.text_input("Personal Email", placeholder="name@example.com")
    with c2:
        position = st.text_input("Position Title", placeholder="Assistant Professor in ...")
        department = st.text_input("College/Department Name", placeholder="College/Department")
        reporting_manager = st.text_input("Reporting Manager’s Title", placeholder="Dean/Chair of ...")
        campus = st.selectbox("Campus", ["Abu Dhabi", "Dubai", "Al Ain"], index=0)
        salary = st.number_input("Total Monthly Compensation (AED)", min_value=0, step=500, value=0)

    st.subheader("Contract Settings")
    c3, c4, c5, c6 = st.columns(4)
    with c3:
        ranks = [k for k in BENEFITS.keys() if k != "_shared"]
        rank = st.selectbox("Rank", ranks, index=ranks.index("Assistant / Lecturer") if "Assistant / Lecturer" in ranks else 0)
    with c4:
        marital_status = st.selectbox("Marital Status", ["Single", "Married"], index=0)
    with c5:
        hire_type = st.selectbox("Hire Type", ["Local", "International"], index=0)
    with c6:
        probation = st.number_input("Probation (months)", min_value=1, max_value=12, value=6, step=1)

    st.divider()
    uploaded_template = st.file_uploader("Upload a custom DOCX template (optional). Leave empty to use the default.", type=["docx"])

    submit = st.form_submit_button("Generate Offer Letter")

if submit:
    today = datetime.now().strftime(DATE_FORMAT)
    base_map = {
        "ID": candidate_id,
        "DATE": today,
        "SALUTATION": salutation,
        "CANDIDATE_NAME": candidate_name,
        "TELEPHONE": telephone,
        "PERSONAL_EMAIL": personal_email,
        "POSITION": position,
        "DEPARTMENT": department,
        "REPORTING_MANAGER": reporting_manager,
        "CAMPUS": campus,
        "SALARY": f"{int(salary):,}" if salary else "",
        "PROBATION": probation,
    }

    benefits_map = compute_benefits_mapping(
        rank=rank,
        marital=marital_status,
        campus=campus,
        is_international=(hire_type == "International"),
    )

    mapping = {**base_map, **benefits_map}
    mapping["BENEFITS_BLOCK"] = compose_benefits_block(mapping)

    required = ["ID", "CANDIDATE_NAME", "PERSONAL_EMAIL", "POSITION", "DEPARTMENT", "REPORTING_MANAGER", "SALARY"]
    missing = [k for k in required if not mapping.get(k)]
    if missing:
        st.warning("Please complete required fields: " + ", ".join(missing))

    try:
        tpl_bytes = uploaded_template.read() if uploaded_template else None
        docx_bytes = generate_docx(tpl_bytes, mapping)
        st.success("Offer letter generated successfully.")
        st.download_button(
            "⬇️ Download Offer Letter (DOCX)",
            data=docx_bytes,
            file_name=f"Offer_{(candidate_name or 'Candidate').replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        st.info(f"Applied rank: {rank} | Marital: {marital_status} | Campus: {campus} | Hire: {hire_type}")
    except Exception as e:
        st.error(f"Generation failed: {e}")

st.caption("Tip: Keep fonts, alignment, and logo in the Word template styles for a perfectly formatted output.")
