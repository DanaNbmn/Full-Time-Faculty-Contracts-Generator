import streamlit as st
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="ADU Faculty Contract Generator", page_icon="📄", layout="centered")

DATE_FORMAT = "%d %B %Y"

# Files bundled with the app
DEFAULT_HEADER_LOGO = "adu_logo.png"     # header logo
DEFAULT_FOOTER_BANNER = "adu_footer.png" # footer banner

# ==== Benefits table (unchanged) ====
BENEFITS = {
    "_shared": {"children_school_allowance": {"AD/Dubai": 60000, "AA": 50000}},
    "Professor": {
        "annual_leave_days": 56,
        "joining_ticket_international": "Economy class air tickets for yourself, your spouse, and up to two (2) eligible dependent children under the age of 21 years residing in the UAE, provided upon commencement of employment.",
        "housing_allowance_k": {"AD/Dubai": {"Single": 45, "Married": 60}, "AA": {"Single": 35, "Married": 45}},
        "furniture_allowance_k_once": {"AD/Dubai": {"Single": 20, "Married": 30}, "AA": {"Single": 20, "Married": 30}},
        "repatriation_allowance": 3000,
    },
    "Associate / Sr. Lecturer": {
        "annual_leave_days": 56,
        "joining_ticket_international": "Economy class air tickets for yourself, your spouse, and up to two (2) eligible dependent children under the age of 21 years residing in the UAE, provided upon commencement of employment.",
        "housing_allowance_k": {"AD/Dubai": {"Single": 45, "Married": 60}, "AA": {"Single": 35, "Married": 45}},
        "furniture_allowance_k_once": {"AD/Dubai": {"Single": 20, "Married": 30}, "AA": {"Single": 20, "Married": 30}},
        "repatriation_allowance": 3000,
    },
    "Assistant / Lecturer": {
        "annual_leave_days": 56,
        "joining_ticket_international": "Economy class air tickets for yourself, your spouse, and up to two (2) eligible dependent children under the age of 21 years residing in the UAE, provided upon commencement of employment.",
        "housing_allowance_k": {"AD/Dubai": {"Single": 45, "Married": 60}, "AA": {"Single": 35, "Married": 45}},
        "furniture_allowance_k_once": {"AD/Dubai": {"Single": 20, "Married": 30}, "AA": {"Single": 20, "Married": 30}},
        "repatriation_allowance": 3000,
    },
    "Senior Instructor": {
        "annual_leave_days": 42,
        "joining_ticket_international": "Economy class air tickets for yourself, your spouse, and up to two (2) eligible dependent children under the age of 21 years residing in the UAE, provided upon commencement of employment.",
        "housing_allowance_k": {"AD/Dubai": {"Single": 35, "Married": 45}, "AA": {"Single": 30, "Married": 40}},
        "furniture_allowance_k_once": {"AD/Dubai": {"Single": 12, "Married": 15}, "AA": {"Single": 12, "Married": 15}},
        "repatriation_allowance": 2000,
    },
    "Instructor": {
        "annual_leave_days": 42,
        "joining_ticket_international": "Economy class air tickets for yourself, your spouse, and up to two (2) eligible dependent children under the age of 21 years residing in the UAE, provided upon commencement of employment.",
        "housing_allowance_k": {"AD/Dubai": {"Single": 35, "Married": 45}, "AA": {"Single": 30, "Married": 40}},
        "furniture_allowance_k_once": {"AD/Dubai": {"Single": 12, "Married": 15}, "AA": {"Single": 12, "Married": 15}},
        "repatriation_allowance": 2000,
    },
}

def campus_key(campus: str) -> str:
    return "AD/Dubai" if campus in ("Abu Dhabi", "Dubai", "AD/Dubai") else "AA"

def fmt_amt(n: int) -> str:
    return f"{int(n):,}"

def compute_benefits(rank: str, marital: str, campus: str, is_international: bool):
    R = BENEFITS[rank]; S = BENEFITS["_shared"]; ckey = campus_key(campus)
    housing = R["housing_allowance_k"][ckey][marital] * 1000
    furniture = R["furniture_allowance_k_once"][ckey][marital] * 1000
    edu = S["children_school_allowance"][ckey]
    return {
        "HOUSING_ALLOWANCE": fmt_amt(housing),
        "FURNITURE_ALLOWANCE": fmt_amt(furniture),
        "JOINING_TICKET": R["joining_ticket_international"] if is_international else "",
        "REPARIATION_ALLOWANCE": fmt_amt(R["repatriation_allowance"]),
        "ANNUAL_LEAVE_DAYS": R["annual_leave_days"],
        "EDUCATION_ALLOWANCE_PER_CHILD": fmt_amt(edu/2),
        "EDUCATION_ALLOWANCE_TOTAL": fmt_amt(edu),
    }

# ---------- helpers ----------
def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")

def bold_prefix_before_colon(paragraph):
    txt = paragraph.text
    if ":" not in txt:
        return
    prefix, rest = txt.split(":", 1)
    prefix += ":"
    for _ in range(len(paragraph.runs)):
        paragraph.runs[0].text = ""
        del paragraph.runs[0]
    r1 = paragraph.add_run(prefix); r1.bold = True; r1.font.size = Pt(11)
    paragraph.add_run(rest)

def apply_header_footer(doc: Document, logo_bytes: bytes | None, footer_bytes: bytes | None):
    """Header/footer on every page; footer spans full width between margins."""
    try:
        doc.settings.odd_and_even_pages_header_footer = False
    except Exception:
        pass

    for section in doc.sections:
        try:
            section.different_first_page_header_footer = False
        except Exception:
            pass

        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

        # ----- Header -----
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try: 
            hp.clear()
        except Exception: 
            pass
        try:
            run = hp.add_run()
            if logo_bytes:
                run.add_picture(BytesIO(logo_bytes), width=Inches(2.0), height=Inches(1.5))
        except Exception:
            pass

        # ----- Footer (span from left margin to right margin) -----
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try: 
            fp.clear()
        except Exception: 
            pass
        try:
            frun = fp.add_run()
            if footer_bytes:
                footer_width = section.page_width - section.left_margin - section.right_margin
                frun.add_picture(BytesIO(footer_bytes), width=footer_width, height=Inches(0.57))
        except Exception:
            pass

def build_letter(m, logo_bytes=None, footer_bytes=None):
    doc = Document()

    doc.add_paragraph(f"Ref: {m['ID']}")
    doc.add_paragraph(f"Date: {m['DATE']}")
    doc.add_paragraph("")

    doc.add_paragraph(f"{m['SALUTATION']} {m['CANDIDATE_NAME']}")
    doc.add_paragraph(f"Tel No: {m['TELEPHONE']}")
    doc.add_paragraph(f"Email ID: {m['PERSONAL_EMAIL']}")
    doc.add_paragraph("")

    doc.add_paragraph(f"Dear {m['SALUTATION']} {m['CANDIDATE_NAME']},")
    doc.add_paragraph(
        f"Abu Dhabi University (ADU) is pleased to offer you a contract of employment for the "
        f"position of {m['POSITION']} in the {m['DEPARTMENT']} based in {m['CAMPUS']}, UAE. "
        f"This position reports to the {m['REPORTING_MANAGER']}."
    )
    doc.add_paragraph(
        "Your first day of employment with Abu Dhabi University will be based on the availability "
        "of legal approvals, and the term of your contract shall be limited to a period of two (2) years, "
        "renewable upon mutual agreement."
    )

    # 1. Package
    add_heading(doc, "1. Package")
    doc.add_paragraph(f"Your total monthly compensation will be AED {m['SALARY']}, comprising:")
    add_bullet(doc, "Basic Salary: 50% of the total monthly compensation.")
    add_bullet(doc, "Other Allowance: 50% of the total monthly compensation.")
    doc.add_paragraph("Payment will be made at the end of each calendar month.")

    # 2. Terms and Conditions
    add_heading(doc, "2. Terms and Conditions")
    doc.add_paragraph(f"Probation Period: The first {m['PROBATION']} months from the start date shall constitute the probationary period.")
    doc.add_paragraph(
        "Notice Period: Upon successful completion of the probationary period, either party may terminate this contract "
        "by providing one academic semester’s written notice, coinciding with the end of the semester, or payment in lieu, "
        "in accordance with ADU policy."
    )

    # 3. Benefits
    add_heading(doc, "3. Benefits")
    doc.add_paragraph(
        f"Accommodation: Unfurnished on-campus accommodation based on availability, or a housing allowance of AED {m['HOUSING_ALLOWANCE']} per year "
        "(paid monthly) will be provided based on eligibility."
    )
    doc.add_paragraph(
        f"Furniture Allowance: AED {m['FURNITURE_ALLOWANCE']} provided at the commencement of employment as a forgivable loan amortized over "
        "three (3) years. Should you leave ADU before completing three years of service, the amount will be repayable on a pro-rata basis."
    )
    doc.add_paragraph(
        "Annual Leave Airfare: Cash in lieu of economy class air tickets for yourself, your spouse, and up to two (2) eligible dependent children "
        "under the age of 21 years residing in the UAE, based on ADU’s published schedule of rates including your country of origin. This amount "
        "will be paid annually in the month of May, prorated to your joining date."
    )
    if m["JOINING_TICKET"]:
        doc.add_paragraph(f"Commencement Air Tickets: {m['JOINING_TICKET']}")
    doc.add_paragraph(
        "Relocation Allowance: Up to AED 3,000 at the commencement of employment to support the relocation of personal effects to ADU-provided accommodation."
        "Reimbursement will be subject to submission of original receipts."
    )
    doc.add_paragraph(
        "Repatriation Air Tickets: You will be provided with Economy Class air tickets for yourself, spouse and your eligible dependents "
        "(up to 2 children under 21 years) residing in the UAE upon your end of employment to your country of origin."
    )
    doc.add_paragraph(
        f"Repatriation Allowance: AED {m['REPARIATION_ALLOWANCE']} upon conclusion of your contract, applicable only upon completion of two (2) "
        "years of continuous service with ADU."
    )
    doc.add_paragraph(
        "Medical Insurance: You will be provided with medical insurance coverage for yourself, spouse and your eligible dependents "
        "(up to 3 children under 21 years) residing in the UAE. (Applicable only for married individuals with spouse/children under their sponsorship)"
    )
    doc.add_paragraph(f"Annual Leave Entitlement: {m['ANNUAL_LEAVE_DAYS']} calendar days of paid annual leave.")
    doc.add_paragraph(
        f"School Fee Subsidy: An annual subsidy of AED {m['EDUCATION_ALLOWANCE_PER_CHILD']} per eligible child under the age of 21 years residing "
        f"in the UAE under your sponsorship, up to a maximum of AED {m['EDUCATION_ALLOWANCE_TOTAL']} per family. This benefit applies only to married "
        "individuals with children under their sponsorship."
    )
    doc.add_paragraph(
        "ADU Tuition Waiver: 75% deduction on tuition fees for self, 50% for dependents and 25% for immediate family in accordance with ADU Policy. "
        "(applicable upon completion of one year of service with ADU)"
    )

    # 4. End of Service Entitlements
    add_heading(doc, "4. End of Service Entitlements")
    doc.add_paragraph(
        "End of Service Gratuity: Calculated at one (1) month’s basic salary for each completed year of service, "
        "in accordance with ADU policy and UAE Labour Law. This will be prorated for any partial years of service. "
        "No gratuity is payable for service of less than one (1) year."
    )

    # 5. Additional Provisions
    add_heading(doc, "5. Additional Provisions")
    doc.add_paragraph(
        "a) If married, any benefits provided by the spouse’s employer will not be duplicated by ADU. You are required to declare to TEG any benefits "
        "at the time of acceptance of this offer and of any future change in your spouse’s benefits while you are employed at ADU."
    )
    doc.add_paragraph("b) You may be requested to teach in other ADU campuses as required by the college.")

    # 6. Documentation Requirements
    add_heading(doc, "6. Documentation Requirements")
    for item in [
        "Ministry of Higher Education and Scientific Research (MOHESR) Clearance.",
        "Completion of visa and ADU sponsorship formalities.",
        "Attestation of original academic qualifications.",
        "Medical examination and clearance.",
        "Notarized copies of marriage and birth certificates (if applicable).",
        "Provision of a current and valid Police/CRB Clearance.",
        "Receipt of satisfactory references.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "By accepting this offer, you attest that all personal and business information and documents provided to ADU are true, accurate, and complete. "
        "You further acknowledge that any discrepancy in such information or documents may lead to the withdrawal of this Employment Offer and the termination "
        "of your employment contract, even after you have joined."
    )

    # 7. Validity
    add_heading(doc, "7. Validity")
    doc.add_paragraph("Your acceptance of this offer must be received within ten (10) working days from the date of this letter.")
    doc.add_paragraph(
        "This offer, once signed, constitutes an official agreement between Abu Dhabi University and yourself. Your signature on this document indicates your agreement "
        "with the included terms and conditions of employment and supersedes all other written and/or verbal agreements, understandings, and offers."
    )
    doc.add_paragraph("We look forward to welcoming you to Abu Dhabi University.")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Prof. Hamad Ebrahim Ali Odhabi")
    doc.add_paragraph("Vice Chancellor for AI and Operational Excellence")

    doc.add_paragraph("")
    add_heading(doc, "Acknowledgment and Acceptance")
    doc.add_paragraph("I accept this offer of employment and agree to sign the Ministry of Labour Contract upon joining ADU.")
    doc.add_paragraph(f"Name (print): {m['SALUTATION']} {m['CANDIDATE_NAME']}")
    doc.add_paragraph("Signature: _______________________________")
    doc.add_paragraph("Date: ___________________________________")

    # Normalize base font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

    # Header & footer on all pages; footer spans margin-to-margin
    apply_header_footer(doc, logo_bytes, footer_bytes)

    # Bold prefixes before ":" across the doc
    for par in doc.paragraphs:
        bold_prefix_before_colon(par)

    out = BytesIO(); doc.save(out); out.seek(0)
    return out.getvalue()

# ---------------- UI ----------------
st.title("📄 ADU – Faculty Contract Generator")

with st.form("offer_form", clear_on_submit=False):
    st.subheader("Candidate & Position")
    c1, c2 = st.columns(2)
    with c1:
        candidate_id = st.text_input("ID (Ref)")
        salutation = st.selectbox("Salutation", ["Dr.", "Mr.", "Ms.", "Prof.", "Eng."])
        candidate_name = st.text_input("Candidate Name")
        telephone = st.text_input("Telephone")
        personal_email = st.text_input("Personal Email")
    with c2:
        position = st.text_input("Position")
        department = st.text_input("Department")
        reporting_manager = st.text_input("Reporting Manager’s Title")
        campus = st.selectbox("Campus", ["Abu Dhabi", "Dubai", "Al Ain"])
        salary = st.number_input("Total Monthly Compensation (AED)", min_value=0, step=500, value=0)

    st.subheader("Contract Settings")
    c3, c4, c5, c6 = st.columns(4)
    with c3:
        rank = st.selectbox("Rank", [k for k in BENEFITS if k != "_shared"], index=2)
    with c4:
        marital_status = st.selectbox("Marital Status", ["Single", "Married"])
    with c5:
        hire_type = st.selectbox("Hire Type", ["Local", "International"])
    with c6:
        probation = st.number_input("Probation (months)", min_value=1, max_value=12, value=6)

    submit = st.form_submit_button("Generate Offer Letter")

if submit:
    today = datetime.now().strftime(DATE_FORMAT)

    base = {
        "ID": candidate_id, "DATE": today, "SALUTATION": salutation, "CANDIDATE_NAME": candidate_name,
        "TELEPHONE": telephone, "PERSONAL_EMAIL": personal_email, "POSITION": position,
        "DEPARTMENT": department, "CAMPUS": campus, "REPORTING_MANAGER": reporting_manager,
        "SALARY": f"{int(salary):,}" if salary else "", "PROBATION": probation,
    }
    benefits = compute_benefits(rank, marital_status, campus, hire_type == "International")
    m = {**base, **benefits}

    # Always use bundled images
    with open(DEFAULT_HEADER_LOGO, "rb") as f:
        logo_bytes = f.read()
    with open(DEFAULT_FOOTER_BANNER, "rb") as f:
        footer_bytes = f.read()

    try:
        docx_bytes = build_letter(m, logo_bytes=logo_bytes, footer_bytes=footer_bytes)
        st.success("Offer letter generated successfully.")
        st.download_button(
            "⬇️ Download Offer Letter (DOCX)",
            data=docx_bytes,
            file_name=f"Offer_{(candidate_name or 'Candidate').replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        st.info(f"Rank: {rank} | Marital: {marital_status} | Campus: {campus} | Hire: {hire_type}")
    except Exception as e:
        st.error(f"Generation failed: {e}")


